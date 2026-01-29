from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
import re
from io import BytesIO
from pdfminer.high_level import extract_text
import tempfile
from docx import Document

ALLOWED = {'.pdf', '.doc', '.docx', '.txt'}
app = Flask(__name__)

# --- SKILL DICTIONARY ---
# A comprehensive list of technical skills to match against case-insensitively
SKILLS_DB = {
    # Languages
    "python", "java", "c++", "c#", "javascript", "typescript", "golang", "ruby", "swift", "kotlin", "php", "rust",
    # Frontend
    "react", "angular", "vue", "next.js", "html", "css", "tailwind", "bootstrap", "redux", "jquery",
    # Backend
    "node.js", "express", "django", "flask", "springboot", "fastapi", "laravel", "asp.net", "graphql", "rest api",
    # Database
    "sql", "mysql", "postgresql", "mongodb", "redis", "firebase", "oracle", "mariadb", "sqlite",
    # Cloud/DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "git", "github", "gitlab", "ci/cd", "terraform", "ansible",
    # Data/AI
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "opencv", "nlp", "tableau", "power bi", "hadoop", "spark"
}

def extract_text_from_file(path, ext):
    """Extract text from various file types"""
    text = ""
    try:
        if ext == '.pdf':
            text = extract_text(path)
        elif ext == '.txt':
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif ext == '.docx':
            doc = Document(path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += '\n' + cell.text
        elif ext == '.doc':
            raise ValueError("Legacy .doc format is not supported. Please convert to .docx or PDF")
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""
    
    return text

def extract_contact_info(text):
    """Extract Email and Phone using Regex"""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    
    # Process phones: filter out short numbers, duplicates
    valid_phones = []
    for p in phones:
        if isinstance(p, tuple): p = "".join(p) # Handle groups
        clean_p = re.sub(r'[^0-9+]', '', p) # Keep only digits and +
        if len(clean_p) >= 10: # Basic length check
            valid_phones.append(p)

    return {
        "email": emails[0] if emails else None,
        "phone": valid_phones[0] if valid_phones else None
    }

def extract_skills(text):
    """Extract skills based on dictionary matching"""
    found_skills = set()
    # Normalize text: lowercase, remove special chars (keep +, . for C++, Node.js)
    # We replace newlines with space to avoid splitting "Node.\njs" issues
    normalized_text = text.lower().replace('\n', ' ')
    
    # Tokenize simply by splitting on non-alphanumeric (but keep + . # for tech names)
    # Actually, simpler approach: Check if skill exists in text as whole word
    # For "C++", simple regex `\b` works poorly with `+`.
    
    for skill in SKILLS_DB:
        # Create a regex to find the skill as a whole word (ish)
        # Escape special chars like +, .
        escaped_skill = re.escape(skill) 
        # Pattern: exact match surrounded by non-word chars OR start/end of string
        # Use simple string check first for speed, then regex for precision if needed?
        # Let's trust simple substring for now, but strict boundary check is better
        
        # Checking logic specific for things like 'c++':
        if skill in ['c++', 'c#', '.net']:
             if skill in normalized_text: # Simple check for special symbols
                 found_skills.add(skill)
        else:
             # Use word boundary for normal words to avoid 'java' matching 'javascript' (handled by order? no)
             # or 'go' matching 'google'
             pattern = r'\b' + escaped_skill + r'\b'
             if re.search(pattern, normalized_text):
                 found_skills.add(skill)

    return list(found_skills)

@app.route("/parse", methods=["POST"])
def parse():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "no file"}), 400
    filename = secure_filename(f.filename)
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED:
        return jsonify({"error": "unsupported"}), 400
    
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    f.save(tmp.name)
    
    try:
        text = extract_text_from_file(tmp.name, ext)
        
        if not text or not text.strip():
            return jsonify({"parsedText": "", "keywords": [], "error": "No text extracted"})
            
        contact_info = extract_contact_info(text)
        skills = extract_skills(text)
        
        return jsonify({
            "parsedText": text,
            "keywords": skills, # Resume Controller expects 'keywords'
            "email": contact_info['email'],
            "phone": contact_info['phone']
        })
        
    except Exception as e:
        return jsonify({"parsedText": "", "keywords": [], "error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except:
            pass

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
