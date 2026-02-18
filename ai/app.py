"""
Flask Application for Resume Parsing and Job Matching.
Exposes endpoints for parsing resumes (PDF/DOCX/TXT) and matching them against job descriptions.
"""
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
import tempfile
from pdfminer.high_level import extract_text
from docx import Document
from resume_parser import parser  # Import NLP Parser
from matcher import JobMatcher # Import Matcher

ALLOWED = {'.pdf', '.doc', '.docx', '.txt'}
app = Flask(__name__)
matcher = JobMatcher()

def extract_text_from_file(path, ext):
    """
    Extracts raw text from a file based on its extension.
    Supports .pdf, .docx, and .txt.
    """
    text = ""
    try:
        if ext == '.pdf':
            text = extract_text(path)
        elif ext == '.txt':
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif ext == '.docx':
            doc = Document(path)
            # Efficiently join paragraphs and tables
            paragraphs = [p.text for p in doc.paragraphs]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tables.append(cell.text)
            text = '\n'.join(paragraphs + tables)
        elif ext == '.doc':
            raise ValueError("Legacy .doc format is not supported. Please convert to .docx or PDF")
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""
    
    return text

@app.route("/parse", methods=["POST"])
def parse():
    """
    Endpoint to parse a resume file.
    Expects a file upload in 'file' field.
    Returns extracted Name, Email, Phone, and Skills.
    """
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "no file"}), 400
    filename = secure_filename(f.filename)
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED:
        return jsonify({"error": "unsupported"}), 400
    
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    f.save(tmp.name)
    
    try:
        text = extract_text_from_file(tmp.name, ext)
        
        if not text or not text.strip():
            return jsonify({"parsedText": "", "keywords": [], "error": "No text extracted"})
            
        # Use NLP Parser to extract entities
        data = parser.parse(text)
        
        return jsonify({
            "parsedText": text,
            "keywords": data["skills"], 
            "name": data["name"],
            "email": data["email"],
            "phone": data["phone"],
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"parsedText": "", "keywords": [], "error": str(e)}), 500
    finally:
        try:
            os.unlink(tmp.name)
        except:
            pass

@app.route("/match", methods=["POST"])
def match():
    """
    Endpoint to calculate match score between resume and job.
    Expects JSON: { "resume_text": str, "job_description": str, "job_keywords": list }
    Returns a match percentage and breakdown.
    """
    data = request.json
    if not data or "resume_text" not in data or "job_keywords" not in data:
        return jsonify({"error": "Missing resume_text or job_keywords"}), 400
    
    resume_text = data["resume_text"]
    job_description = data.get("job_description", "")
    job_keywords = data["job_keywords"]
    
    try:
        result = matcher.match(resume_text, job_description, job_keywords)
        return jsonify(result)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
